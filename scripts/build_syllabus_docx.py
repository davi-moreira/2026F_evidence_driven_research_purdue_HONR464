#!/usr/bin/env python3
"""Build the official HONR 46400 syllabus .docx.

Format is cloned from the QM 47400 Fall 2026 syllabus (Davi's hand-authored
Word document): that file is opened as the base, so styles.xml, numbering.xml,
the theme, page setup, and the portrait + landscape section pair all carry over
untouched. Only the body is rebuilt, and only the header banner is swapped (the
Mitch Daniels lockup is cropped to the Purdue wordmark, since HONR 46400 is a
John Martinson Honors College course).

Content comes from this repository, never retyped:
  * syllabus.qmd                     - the assessment contract and policies
  * planning/MEETING_SCHEDULE.csv    - the 43-meeting backbone
  * scripts/notebooks_map.py         - the notebook <-> meeting mapping
  * brightspace/gradebook_spec.md    - milestone due dates

Run:  .venv/bin/python scripts/build_syllabus_docx.py
"""
from __future__ import annotations

import csv
import re
import shutil
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document                                    # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                # noqa: E402
from docx.oxml.ns import qn                                  # noqa: E402
from docx.shared import Inches, Pt                           # noqa: E402
from notebooks_map import REPO_SLUG, nb_of, student_filename  # noqa: E402
from validate_calendar import no_class_days                  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
QM_DOCX = (ROOT.parent.parent / "predictive_analytics"
           / "2026F_predictive_analytics_QM474" / "_syllabus" / "2026F"
           / "2026F_predictive_analytics_purdue_QM474.docx")
OUT = ROOT / "_syllabus" / "2026F" / "2026F_evidence_driven_research_purdue_HONR464.docx"
CSV_PATH = ROOT / "planning" / "MEETING_SCHEDULE.csv"
SRC_LOGO = ROOT / "_syllabus" / "complementary_material" / "logo.png"
BANNER = ROOT / "_syllabus" / "complementary_material" / "purdue_wordmark_banner.png"
# logo.png is the Purdue | Mitch Daniels School of Business lockup, 1423x152.
# HONR 46400 is a John Martinson Honors College course, so the Daniels half is
# cropped off at the divider bar (columns 799-802) and only the Purdue
# University wordmark is kept.
BANNER_CROP_W = 767

SITE = f"https://davi-moreira.github.io/{REPO_SLUG.split('/')[-1]}/"
# Davi's Outlook "Book with me" page, as used in the QM474 syllabus.
BOOKING_URL = ("https://outlook.office.com/bookwithme/user/"
               "cf2ec4ce741a4d7aa70889bd35d00e55@purdue.edu/meetingtype/"
               "xG_mY-1o4EeSutf45nmMQg2?anonymous&ep=mcard")
POSTERS_URL = "https://davi-moreira.github.io/applied_projects.html"

# ── logistics (settled 2026-07-31, confirmed by Davi 2026-08-23) ─────────────
CLASS_LINE = ("Section 002 (CRN 25596): Monday, Wednesday, Friday, "
              "1:30–2:20 p.m., HCRS-1054.")
OFFICE = "Young Hall 1019"
OFFICE_HOURS = ("Wednesdays: 2:30 p.m.–3:30 p.m. (EST), in person, "
                "HCRN 1095, or by appointment.")

# The live chain, read from the ONE declaration. Duplicating it here is how a
# retired milestone (D54) or a moved deadline (D55) silently survives in a
# generated .docx after the source of truth has moved on.
def _milestone_due() -> dict[int, str]:
    import yaml
    cfg = yaml.safe_load((ROOT / "course_config.yaml").read_text())
    return {int(k[1:]): v["due"] for k, v in cfg["milestones"].items()}


MILESTONE_DUE = _milestone_due()


# ── low-level docx helpers ──────────────────────────────────────────────────
def clear_body(doc: Document):
    """Empty the body but keep BOTH section definitions.

    QM474's document is two sections: a portrait front matter and a landscape
    schedule. The portrait sectPr lives inside a paragraph's pPr (that is how
    Word stores a section break); the landscape one sits at body level. Deleting
    every paragraph would silently destroy the portrait section and render the
    whole syllabus landscape, so lift the portrait sectPr out first and hand it
    back for reinsertion at the section break.
    """
    import copy
    body = doc.element.body
    portrait = None
    for sp in body.iter(qn("w:sectPr")):
        if sp.getparent().tag.endswith("}pPr"):
            portrait = copy.deepcopy(sp)
            break
    if portrait is None:
        raise SystemExit("✗ portrait sectPr not found in the format base")
    for child in list(body.iterchildren()):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)
    return portrait


def end_portrait_section(doc, portrait_sectpr):
    """Close the portrait section: an empty paragraph whose pPr carries it."""
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().append(portrait_sectpr)
    return p


# QM474's lists are real Word numbered lists. These are the numId values its
# numbering.xml already defines; reusing them keeps the numbering identical.
NUM = {
    "outcomes": 3,       # LEARNING OUTCOMES
    "criteria": 6,       # Final Project criteria
    "netiquette": 8,
    # numId 2 is a Symbol bullet indented 1440 twips; the AI-policy list sits at
    # 720, matching the list Word created when Davi bulleted this section by
    # hand. numId 14 is that definition and is otherwise unused in this build.
    "ai_policy": 14,
    "challenge": 4,      # ilvl 0 for the two steps, ilvl 1 for the fields
    "extra_credit": 5,
    "caps": 1,           # ilvl 3
}

# "Body Text" is defined at 24pt in this document; QM474 overrides every run to
# 12pt. Match that or the paragraph renders enormous.
BODY_TEXT_PT = Pt(12)


def number(p, num_id: int, ilvl: int = 0):
    """Attach an existing numbering definition to a paragraph."""
    pPr = p._p.get_or_add_pPr()
    numpr = pPr.makeelement(qn("w:numPr"), {})
    lvl = pPr.makeelement(qn("w:ilvl"), {qn("w:val"): str(ilvl)})
    nid = pPr.makeelement(qn("w:numId"), {qn("w:val"): str(num_id)})
    numpr.append(lvl)
    numpr.append(nid)
    pPr.append(numpr)
    return p


# ── Davi's Word pass, folded back (D62, extended by D74) ────────────────────
# He edits the BUILT .docx in Word, and his 2026-08-23 pass added blank
# paragraphs between the body blocks so the printed syllabus breathes. A second
# pass on 2026-08-25 tightened two of those runs from 3 to 1 and reworked the
# schedule header; it was captured on 2026-08-31 under D74, after the .docx was
# found newer than this script. The spacing is his instruction, so the generator
# produces it. Each entry is
# (paragraph-text prefix, how many empty paragraphs must sit immediately above
# it). Applied as a post-pass so the content code below stays readable.
BLANK_ABOVE = [
    ("COURSE WEBSITE:", 1),
    ("COURSE OBJECTIVES:", 1),
    ("LEARNING OUTCOMES:", 1),
    ("COURSE MATERIALS:", 1),
    ("Course book (free, required):", 1),
    ("Recommended: Blair", 1),
    ("Computing (Required):", 1),
    ("AI tools (Required):", 1),
    ("ASSESSMENTS", 1),
    ("The course assesses the research chain", 1),
    ("Grading: Final letter grades", 1),
    ("Grading scale:", 1),   # D74: Davi's Word pass tightened 3 -> 1
    ("Attendance:", 1),
    ("Participation:", 1),
    ("IYT Practice:", 1),
    ("Lecture Notebooks:", 1),
    ("EXTRA CREDIT OPPORTUNITIES:", 1),
    ("Course Evaluations:", 1),
    ("Issues in the Course Materials:", 1),
    ("AI POLICY:", 1),       # D74: Davi's Word pass tightened 3 -> 1
    ("AI is a research assistant", 1),
    ("Some activities may be done without AI", 1),
    ("Deadlines: The Registrar", 1),
    ("Academic Integrity:", 1),
    ("Accessibility and Accommodations:", 1),
    ("CAPS Information:", 1),
    ("Non-Discrimination Statement:", 1),
    ("Basic Needs Security:", 1),
    ("Emergency Situation(s):", 1),
    ("Subject to Change Policy:", 1),
    ("Tentative Course Schedule", 1),
]


def apply_blank_spacing(doc):
    """Make each BLANK_ABOVE block sit under exactly N empty paragraphs."""
    from copy import deepcopy
    for prefix, want in BLANK_ABOVE:
        target = next((p for p in doc.paragraphs
                       if p.text.strip().startswith(prefix)), None)
        if target is None:
            raise SystemExit(f"✗ blank-spacing anchor not found: {prefix!r}")
        el = target._p
        have = 0
        probe = el.getprevious()
        while probe is not None and probe.tag == el.tag and not "".join(
                probe.itertext()).strip():
            have += 1
            probe = probe.getprevious()
        for _ in range(want - have):
            blank = deepcopy(el)
            for child in list(blank):
                if not child.tag.endswith("}pPr"):
                    blank.remove(child)
            el.addprevious(blank)


def para(doc, text="", style="Normal", bold_prefix=None, space_after=None,
         align=None):
    """A paragraph, optionally opening with a bold run-in label (QM474's idiom)."""
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    if text:
        p.add_run(text)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if style == "Body Text":
        for r in p.runs:
            r.font.size = BODY_TEXT_PT
    return p


def caps_heading(doc, label, text=""):
    """QM474 writes section heads as an ALL-CAPS bold run-in label, not a Word
    heading. Reproduce that exactly."""
    return para(doc, text, style="Normal", bold_prefix=label)


def bullet(doc, text, bold_prefix=None):
    return para(doc, text, style="List Paragraph", bold_prefix=bold_prefix)


def clone_table_format(src_tbl, dst_tbl):
    """Graft QM474's own tblPr and tblGrid onto our table.

    Setting a style name and per-cell widths through python-docx is not enough:
    the borders live in the source table's tblPr and the column proportions in
    its tblGrid, and an autofit table silently discards assigned widths. Copying
    both elements reproduces QM474's layout exactly, so long as the column count
    matches.
    """
    import copy
    src_cols = len(src_tbl.columns)
    dst_cols = len(dst_tbl.columns)
    if src_cols != dst_cols:
        raise SystemExit(f"✗ column mismatch: source {src_cols}, target {dst_cols}")
    for tag in ("w:tblPr", "w:tblGrid"):
        old = dst_tbl._tbl.find(qn(tag))
        if old is not None:
            dst_tbl._tbl.remove(old)
    grid = copy.deepcopy(src_tbl._tbl.find(qn("w:tblGrid")))
    pr = copy.deepcopy(src_tbl._tbl.find(qn("w:tblPr")))
    dst_tbl._tbl.insert(0, grid)
    dst_tbl._tbl.insert(0, pr)
    # carry the source's per-column cell widths onto every row
    widths = [c._tc.find(qn("w:tcPr")).find(qn("w:tcW"))
              for c in src_tbl.rows[0].cells]
    for row in dst_tbl.rows:
        for cell, w in zip(row.cells, widths):
            if w is None:
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            old = tcPr.find(qn("w:tcW"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(copy.deepcopy(w))


def hyperlink(paragraph, url, text, bold=False):
    """python-docx has no hyperlink API; build the w:hyperlink element."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = paragraph._p.makeelement(qn("w:r"), {})
    rpr = paragraph._p.makeelement(qn("w:rPr"), {})
    style = paragraph._p.makeelement(qn("w:rStyle"), {qn("w:val"): "Hyperlink"})
    rpr.append(style)
    if bold:
        rpr.append(paragraph._p.makeelement(qn("w:b"), {}))
    run.append(rpr)
    t = paragraph._p.makeelement(qn("w:t"), {})
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def prune_orphan_hyperlinks(doc):
    """Drop hyperlink relationships inherited from the format base.

    The document is built by copying QM474's file and emptying its body, so its
    relationship part still lists every URL QM474 linked - Kaggle, statlearning,
    the QM474 course site, the honors-contract pages, and QM474's own Colab
    notebooks. Nothing references them any more, but they travel inside the
    .docx and would show a reader another course's links. Remove the external
    hyperlink rels that no run points at, and leave every other relationship
    (styles, numbering, header, footer, theme, images) untouched.
    """
    used = {h.get(qn("r:id")) for h in doc.element.body.iter(qn("w:hyperlink"))}
    used.discard(None)
    HYPERLINK = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                 "relationships/hyperlink")
    dropped = 0
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == HYPERLINK and rel.is_external and rid not in used:
            doc.part.rels.pop(rid)
            dropped += 1
    return dropped


def patch_header_banner(docx_path: Path, image_path: Path, ratio: float):
    """Swap the header banner in the saved package.

    python-docx will not write a replacement image blob back through the
    relationship, and edits to the header part's drawing extents do not survive
    the save either, so do both directly on the .docx zip after saving: replace
    word/media/image1.png and scale every cx in word/header1.xml by the crop
    ratio, so the wordmark prints at exactly the size it did in QM474's header.
    """
    import zipfile
    new_png = image_path.read_bytes()

    def scale(m):
        return f'{m.group(1)}cx="{int(int(m.group(2)) * ratio)}"'

    src = zipfile.ZipFile(docx_path)
    items = [(i, src.read(i.filename)) for i in src.infolist()]
    src.close()
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as out:
        for info, data in items:
            if info.filename == "word/media/image1.png":
                data = new_png
            elif info.filename == "word/header1.xml":
                xml = data.decode("utf-8")
                xml = re.sub(r'(<wp:extent )cx="(\d+)"', scale, xml)
                xml = re.sub(r'(<a:ext )cx="(\d+)"', scale, xml)
                data = xml.encode("utf-8")
            out.writestr(info, data)


# ── schedule ────────────────────────────────────────────────────────────────
def week_of(unit: str) -> str:
    m = re.match(r"Week (\d+)", unit or "")
    return m.group(1) if m else ""


def fmt_date(iso: str) -> str:
    d = date.fromisoformat(iso)
    return f"{d.strftime('%a')} {d.strftime('%b')} {d.day}"


# Milestone notes are derived from the authoritative due-date map, not parsed
# out of the planning prose. The prose is not safely parseable: several
# milestone_developed cells carry a semicolon *inside* the parentheses, which
# silently splits M13 away from its own deadline.
# Deadlines that are NOT 11:59 PM on their own date, and the terminal one.
CLOCK_NOTE = {11: "at class", 12: "2:30 PM"}
TERMINAL = {13: "TERMINAL, no revision window, submitted for printing"}

# D55 put almost every milestone on a Sunday, which is never a class meeting.
# A milestone is therefore announced on its own date when that date IS a
# meeting, and otherwise on the last meeting before it. Nothing is hardcoded:
# drop a milestone from course_config.yaml and its announcement disappears with
# it, exactly like the handout PDFs and the schedule badges.
def _announcement_dates() -> dict[str, list[str]]:
    meetings = sorted({r["date"] for r in
                       csv.DictReader(open(CSV_PATH, newline=""))})
    out: dict[str, list[str]] = {}
    for n in sorted(MILESTONE_DUE):
        due = MILESTONE_DUE[n]
        d = date.fromisoformat(due)
        label = f"M{n} DUE"
        if due in meetings:
            when = due
            if n in CLOCK_NOTE:
                label += f" {CLOCK_NOTE[n]}"
        else:
            earlier = [m for m in meetings if m < due]
            if not earlier:
                continue
            when = earlier[-1]
            label = (f"M{n} due {d.strftime('%a')} "
                     f"{d.strftime('%b')} {d.day}, 11:59 PM")
        if n in TERMINAL:
            label += f" — {TERMINAL[n]}"
        out.setdefault(when, []).append(label)
    return out


_ANNOUNCE: dict[str, list[str]] | None = None


def milestone_note(iso: str) -> str:
    global _ANNOUNCE
    if _ANNOUNCE is None:
        _ANNOUNCE = _announcement_dates()
    return "; ".join(_ANNOUNCE.get(iso, []))


def build_schedule_rows():
    rows = list(csv.DictReader(open(CSV_PATH, newline="")))
    tracked = {p.name for p in (ROOT / "notebooks" / "student").glob("*.ipynb")}
    out = []
    for r in rows:
        n = nb_of(r["other_material"])
        nb_label, nb_url = "—", None
        if n is not None and student_filename(n) in tracked:
            nb_label = f"nb{n:02d} Open in Colab"
            nb_url = (f"https://colab.research.google.com/github/{REPO_SLUG}/"
                      f"blob/main/notebooks/student/{student_filename(n)}")
        # The two asynchronous November sessions are announced separately
        # (_announcements/), so the printed schedule reads like every other row.
        title = r["title"].replace("ASYNC — ", "")
        out.append({
            "wk": week_of(r["unit"]),
            "date": fmt_date(r["date"]),
            "topic": title,
            "nb": nb_label,
            "nb_url": nb_url,
            "note": milestone_note(r["date"]),
            "iso": r["date"],
        })
    # the Expo is a required meeting outside the MWF pattern; it is not a
    # numbered class meeting, so the CSV does not carry it.
    expo = {
        "wk": "13", "date": "Tue Nov 17",
        "topic": "PURDUE FALL UNDERGRADUATE RESEARCH EXPO — required poster presentation",
        "nb": "—", "nb_url": None,
        "note": "POSTER PRESENTATION (graded in the Final Project)",
        "iso": "2026-11-17",
    }
    out.append(expo)

    # The five MWF days the term does not meet. They carry no meeting number, so
    # the CSV has no row for them, and without a row the printed schedule jumps
    # from one meeting to the next with no sign that class was off. The Wk label
    # comes from a meeting in the SAME calendar week.
    for iso, label in no_class_days().items():
        d = date.fromisoformat(iso)
        wk = next((r["wk"] for r in out
                   if date.fromisoformat(r["iso"]).isocalendar()[:2]
                   == d.isocalendar()[:2] and r["wk"]), "")
        out.append({"wk": wk, "date": fmt_date(iso), "topic": label,
                    "nb": "—", "nb_url": None, "note": "", "iso": iso,
                    "noclass": True})

    out.sort(key=lambda r: (r["iso"], r["topic"].startswith("PURDUE")))
    return out


# ── the document ────────────────────────────────────────────────────────────
def main():
    if not QM_DOCX.exists():
        raise SystemExit(f"✗ QM474 format base not found: {QM_DOCX}")
    if not BANNER.exists():
        if not SRC_LOGO.exists():
            raise SystemExit(f"✗ neither {BANNER.name} nor {SRC_LOGO.name} found "
                             f"in {SRC_LOGO.parent}")
        from PIL import Image
        src = Image.open(SRC_LOGO).convert("RGBA")
        src.crop((0, 0, BANNER_CROP_W, src.height)).save(BANNER)
        print(f"  cropped {SRC_LOGO.name} -> {BANNER.name} "
              f"(Purdue wordmark only)")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(QM_DOCX, OUT)
    doc = Document(OUT)
    fmt = Document(QM_DOCX)      # read-only: the source of table formatting
    portrait_sectpr = clear_body(doc)

    # ── title block ─────────────────────────────────────────────────────────
    t = para(doc, "HONR 46400: Evidence-Driven Research", space_after=0)
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)
    s = para(doc, "Fall 2026", space_after=12)
    s.runs[0].bold = True
    s.runs[0].font.size = Pt(13)

    para(doc, "\tDavi Moreira", bold_prefix="Instructor:", space_after=0)
    para(doc, f"\t{OFFICE}", bold_prefix="Office:", space_after=0)
    para(doc, "\tdcordeir@purdue.edu", bold_prefix="Email:", space_after=6)

    para(doc, f"\t{CLASS_LINE}", bold_prefix="Class:", space_after=6)

    para(doc, "", bold_prefix="Office Hours:", space_after=0)
    para(doc, OFFICE_HOURS, space_after=6)

    para(doc, "Individual Appointments (Zoom):", space_after=0)
    p = para(doc, "", space_after=6)
    p.add_run("Book time with me: ")
    hyperlink(p, BOOKING_URL, "Virtual - General Office Hours")
    p.add_run(", or by appointment.")

    para(doc, "Email responses are typically within 24 business hours. If for "
              "some reason you do not have a response by the 24-hour mark, "
              "please email me again.", bold_prefix="Note: ", space_after=10)

    # ── course infrastructure ───────────────────────────────────────────────
    p = caps_heading(doc, "COURSE BRIGHTSPACE: ")
    hyperlink(p, "https://purdue.brightspace.com/", "Purdue Brightspace")
    p.add_run(". The official material, information, and important "
              "announcements about the course will be posted on Brightspace. "
              "You should check the course page on a regular basis.")

    p = caps_heading(doc, "COURSE WEBSITE: ",
                     "This is an additional, non-official, source of material "
                     "for our course:")
    p = para(doc, "")
    hyperlink(p, SITE, SITE)

    caps_heading(
        doc, "COURSE OBJECTIVES:",
        " In this course, Honors students learn how to turn curiosity into "
        "credible, evidence-based insight without requiring a strong background "
        "in quantitative methods or computing. Students will learn to identify "
        "meaningful gaps in a chosen field, translate those gaps into "
        "well-scoped research problems, and formulate research questions that "
        "can be answered with data. Through a structured, supported workflow, "
        "students learn how to select and justify an appropriate quantitative "
        "method approach (description, statistical inference, predictive "
        "modeling, and/or causal reasoning) and how to use computational tools "
        "and AI to locate sources, operationalize concepts, analyze evidence, "
        "verify results, and document decisions responsibly. The course's "
        "defining message is that AI is your arm and your research assistant, "
        "not your brain.")

    para(doc, " By the conclusion of this course, students will be able to:",
         style="Body Text", bold_prefix="LEARNING OUTCOMES:")
    for item in [
        "Identify meaningful gaps in a chosen field and translate them into "
        "well-scoped research problems.",
        "Formulate research questions that can be answered with data.",
        "Select and justify an appropriate quantitative method approach: "
        "description, statistical inference, predictive modeling, and/or "
        "causal reasoning.",
        "Use computational tools and AI to locate sources, operationalize "
        "concepts, analyze evidence, verify results, and document decisions "
        "responsibly.",
        "Design and defend a rigorous research approach.",
        "Interpret findings with appropriate uncertainty and limitations.",
        "Communicate results effectively in written and oral formats.",
    ]:
        number(para(doc, item, style="Normal (Web)"), NUM["outcomes"])

    caps_heading(doc, "COURSE MATERIALS:")
    p = para(doc, "", style="Normal")
    r = p.add_run("Course book (free, required): ")
    r.bold = True
    hyperlink(p, SITE + "book/",
              "EDR|AI — Evidence-Driven Research in the Age of AI: How to "
              "Design, Analyze, Verify, and Defend")
    p.add_run(". It walks through the material at an undergraduate reading "
              "level, and every chapter has its own companion Colab notebook "
              "where you run the chapter's code and complete its closing "
              "“It is your turn” section. EDR|AI is a work in "
              "progress: it is still under development and will keep growing "
              "and improving during the semester.")

    p = para(doc, "", style="Normal")
    r = p.add_run("Recommended: ")
    r.bold = True
    p.add_run("Blair, G., Coppock, A., & Humphreys, M. (2023). Research Design "
              "in the Social Sciences: Declaration, Diagnosis, and Redesign. "
              "Princeton University Press. Read free online at ")
    hyperlink(p, "https://book.declaredesign.org/", "book.declaredesign.org")
    p.add_run(". Each week the course-book chapters are required and the "
              "matching chapters of this text are recommended.")

    para(doc, " A laptop or desktop with internet access and a modern web "
              "browser. Course activities run in the browser through Google "
              "Colab (https://colab.research.google.com/), with no local "
              "installation and no purchases required.",
         bold_prefix="Computing (Required):")

    para(doc, " The course teaches responsible, documented AI use (see the AI "
              "Policy below). Your own AI assistant is the primary in-notebook "
              "research partner; Purdue GenAI Studio supplies reviewer roles; "
              "Microsoft Copilot and other general-purpose tools are permitted "
              "with disclosure.", bold_prefix="AI tools (Required):")
    # ── assessments ─────────────────────────────────────────────────────────
    caps_heading(doc, "ASSESSMENTS")
    para(doc, "The course assesses the research chain itself: a semester-long "
              "research project carried through milestones, presented publicly "
              "at the Purdue Undergraduate Research Conference, and closed with "
              "a reproducible package, a bounded research note, and a written "
              "reflection on the semester's work.")
    # The scale line below the table already ends "(no curve)", so this lead-in
    # does not repeat it (Davi, 2026-08-23 Word pass).
    para(doc, "Final letter grades follow the fixed scale below.",
         bold_prefix="Grading: ")

    weights = [("Assessment", "Weight"), ("Attendance", "1%"),
               ("Participation", "9%"), ("IYT Practice", "15%"),
               ("Lecture Notebooks", "20%"),
               ("Final Project", "55%"), ("Total", "100%")]
    tw = doc.add_table(rows=len(weights), cols=2)
    tw.style = doc.styles["Grid Table 1 Light"]
    clone_table_format(fmt.tables[0], tw)
    for i, (a, b) in enumerate(weights):
        tw.rows[i].cells[0].text = a
        tw.rows[i].cells[1].text = b
        if i == 0 or a == "Total":
            for c in tw.rows[i].cells:
                for r in c.paragraphs[0].runs:
                    r.bold = True
    para(doc, "")
    # The scale belongs with the weights it grades, not at the end of the Final
    # Project criteria where QM474 keeps it (Davi, 2026-08-23 Word pass).
    para(doc, "A ≥ 93, A− ≥ 90, B+ ≥ 87, B ≥ 83, B− ≥ 80, C+ ≥ 77, C ≥ 73, "
              "C− ≥ 70, D ≥ 60, F < 60 (no curve).",
         bold_prefix="Grading scale: ")

    para(doc, " Attendance is checked with iClicker. A minimum attendance rate "
              "of 85% is required to meet the expectations of this course, "
              "because the course depends on live investigation, feedback, and "
              "project work that cannot be fully reconstructed afterward.",
         bold_prefix="Attendance:")
    para(doc, " Participation is graded as one block. It covers the required "
              "course surveys, the course-closing reflection, and other "
              "constructive contributions to the course. These items are scored "
              "for completion and timeliness, never for the opinions they "
              "express, and the lowest few scores are dropped automatically. "
              "The full list, the due dates, and the submission instructions "
              "are posted on the course page.",
         bold_prefix="Participation:")
    para(doc, " Each required EDR|AI chapter closes with an \u201cIt is your "
              "turn\u201d section that you complete in that chapter's companion "
              "Colab notebook. Each section is due at 11:59 PM on the date that "
              "chapter's reading was due. This work is graded for completion, not "
              "for the conclusions you reach: full credit when it arrives on time, "
              "half credit when it arrives within seven days of the deadline. The "
              "lowest few credits are dropped automatically. The dated list lives "
              "on the course schedule page, where each meeting row marks the "
              "sections due that night, and on Brightspace, where you submit them.",
         bold_prefix="IYT Practice:")
    para(doc, " Each week you work in that week's lecture notebook during "
              "class, and you hand that same notebook in at the end of the week. "
              "This work is graded for completion, not for whether your answers "
              "came out right and not for anything you say in the lab meeting "
              "that opens class: full credit when it arrives on time, half "
              "credit when it arrives within seven days of the deadline. The "
              "lowest few credits are dropped automatically. The due dates and "
              "the submission instructions are posted on the course page.",
         bold_prefix="Lecture Notebooks:")

    # the Final Project section — QM474's text, the D53 wording, verbatim
    para(doc, " Students will complete a practical evidence-driven research "
              "project culminating in a poster presentation at the "
              "Undergraduate Research Conference. Individual work is the "
              "default; you may instead complete the project in a group with "
              "the instructor's approval before shared work begins. A "
              "comprehensive set of project guidelines will be provided, and "
              "the assessment structure will adhere to the following criteria, "
              "whose percentages are shares of your final course grade and "
              "together make up the Final Project's 55%:",
         style="Normal (Web)", bold_prefix="Final Project:")
    for label, body in [
        ("Milestone Deliverables (20%): ",
         "Students will submit incremental project components on specific due "
         "dates. These deliverables allow for early feedback and ensure steady "
         "progress throughout the semester. Grades will reflect each "
         "milestone's clarity, completeness, and timely submission."),
        ("Peer Evaluation (10%): ",
         "To encourage accountability and productive research work, students "
         "will evaluate their peers' contributions. These assessments help "
         "ensure balanced participation and measure collaborative "
         "effectiveness."),
        ("Peer Review (5%): ",
         "Each student will review and provide constructive feedback on the "
         "other projects' posters. This process encourages engagement, "
         "enhances critical analysis skills, and promotes a culture of "
         "constructive critique."),
        ("Poster Presentation at the Purdue Undergraduate Research Conference "
         "(10%): ",
         "A poster template and assessment rubric will be shared, and you are "
         "encouraged to review previous award-winning student posters on my "
         f"personal website {POSTERS_URL} for "
         "inspiration. Your final posters must be submitted by the due date "
         "indicated in the course schedule, after which they will be printed "
         "and distributed during a dedicated Poster Presentation Preparation "
         "class. Additional details on the conference can be found at "
         "https://www.purdue.edu/undergrad-research/conferences/index.php. As "
         "the event may not coincide with our regular class time, please "
         "communicate with your other course instructors in advance regarding "
         "potential scheduling conflicts. If any issues arise, please let me "
         "know. We will not hold our usual class immediately following the "
         "Poster Presentation, allowing you time to rest and catch up on other "
         "coursework. Consult the course schedule for further details."),
        ("Instructor/TA Evaluation (10%): ",
         "After the Undergraduate Research Conference your instructor will "
         "evaluate your final submission based on a rubric that will be "
         "shared."),
    ]:
        number(para(doc, body, style="Normal (Web)", bold_prefix=label),
               NUM["criteria"])

    # ── grade challenges ────────────────────────────────────────────────────
    para(doc, " Students who wish to challenge a grade must do so within 3 "
              "calendar days of the grade's release. An exception applies for "
              "the last week of class, when assignments must be challenged "
              "within 1 calendar day to ensure timely computation of final "
              "course grades. Grade challenges must be grounded in legitimate "
              "disputes over research-methods principles or grading accuracy. "
              "Challenges based on post-hoc legalistic arguments or subjective "
              "dissatisfaction will not be considered.",
         bold_prefix="Grade Challenges:")
    para(doc, "To challenge an assignment score, follow these steps:")
    for item, lvl in [
        ("Review the posted solutions or rubrics thoroughly.", 0),
        ("If you still believe an error has been made, email the instructor "
         "with the following information:", 0),
        ("Your course name, section, and lecture date/time.", 1),
        ("Your name and student ID.", 1),
        ("Assignment number/title.", 1),
        ("Specific deduction being challenged.", 1),
        ("Reason for the challenge, clearly explaining why the deduction is "
         "believed to be incorrect, with reference to the solutions or grading "
         "rubric.", 1),
    ]:
        number(bullet(doc, item), NUM["challenge"], lvl)
    para(doc, "Note that grades will not be discussed in the classroom at any "
              "time—before, during, or after class. Students are encouraged to "
              "attend office hours for questions related to course content or "
              "assignment clarification. After the challenge period, all grades "
              "are final and cannot be revised further for purposes of "
              "calculating final course grades.")

    # ── extra credit (Davi 2026-08-23: keep these two only) ─────────────────
    caps_heading(doc, "EXTRA CREDIT OPPORTUNITIES:")
    number(bullet(doc, "You will be provided with an opportunity to anonymously "
                "evaluate this course and your instructor. You will receive an "
                "official email from evaluation administrators with a link to "
                "the online evaluation site. Your feedback is incredibly "
                "valuable. I strongly urge you to participate in the evaluation "
                "system. To encourage your participation, each student who "
                "submits a screenshot of their final course evaluation "
                "submission confirmation will receive a 0.5% bonus, which will "
                "be added to their final Participation grade.",
           bold_prefix="Course Evaluations: "), NUM["extra_credit"])
    number(bullet(doc, "Students can earn 0.25% extra credit towards their final "
                "Participation points for each valid issue they identify in "
                "course materials and report via email by Friday, December 4, "
                "at 11:59 pm. If the issue is confirmed, I will acknowledge it "
                "in my reply. To receive credit, students must upload the PDF "
                "containing their original email notifications along with my "
                "responses confirming the issues in the last week of the "
                "course.",
           bold_prefix="Issues in the Course Materials: "), NUM["extra_credit"])

    # ── netiquette + AI policy + university policies ────────────────────────
    para(doc, "NETIQUETTE GUIDELINES (for Zoom Classes or Office Hours):",
         style="Normal (Web)")
    for item in [
        "Be respectful and clear in communication; contribute constructively "
        "to in-class discussion and peer critique.",
        "In any online session, join with your full name and mute when not "
        "speaking.",
        "Avoid background distractions.",
    ]:
        number(para(doc, item, style="Normal (Web)"), NUM["netiquette"])

    caps_heading(doc, "AI POLICY:")
    # Davi's 2026-08-23 Word pass broke the policy body out of running prose
    # into one bullet per principle, so each rule can be pointed at on its own.
    # The lead-in and the closing no-AI sentence stay plain paragraphs.
    para(doc, "AI is a research assistant in this course, not a substitute for "
              "your judgment. You may use approved AI tools for course work "
              "unless an activity explicitly says otherwise. This course "
              "teaches responsible, documented AI use:", style="Body Text")
    for principle in [
        "AI is your arm and your research assistant, not your brain.",
        "Follow three rules: ask clearly, verify the output, and document "
        "what you used and decided.",
        "Record required uses in your AI Research Ledger.",
        "Check every factual claim, citation, source, calculation, and result "
        "against real evidence.",
        "You remain responsible for everything you submit and must be able to "
        "explain and defend it.",
        "Research questions, design choices, measurement decisions, ethical "
        "judgments, claims, uncertainty, and limitations remain your decisions.",
        "Undisclosed or unverified AI output may be treated as an "
        "academic-integrity violation.",
    ]:
        number(para(doc, principle, style="Body Text"), NUM["ai_policy"])
    para(doc, "Some activities may be done without AI. Any no-AI activity will "
              "be clearly identified in its instructions.",
         style="Body Text")

    para(doc, " The Registrar's add, drop, and modify deadline dates for this "
              "semester are at the following website: "
              "http://www.purdue.edu/Registrar/", bold_prefix="Deadlines:")
    para(doc, "  Please refer to the University Policies on Brightspace.",
         bold_prefix="Academic Integrity:")
    para(doc, "Purdue University strives to make learning experiences as "
              "accessible as possible. If you anticipate or experience physical "
              "or academic barriers based on disability, please let me know so "
              "we can discuss options. You are also encouraged to contact the "
              "Disability Resource Center: drc@purdue.edu or 765-494-1247.",
         bold_prefix="Accessibility and Accommodations: ")
    number(bullet(doc, "Purdue University is committed to advancing the mental health "
                "and well-being of its students. If you or someone you know is "
                "feeling overwhelmed, depressed, or in need of support, "
                "services are available. Contact Counseling and Psychological "
                "Services (CAPS) at 765-494-6995 and http://www.purdue.edu/caps/ "
                "during and after hours, on weekends and holidays, or through "
                "counselors located in the Purdue University Student Health "
                "Center (PUSH) during business hours.",
           bold_prefix="CAPS Information: "), NUM["caps"], 3)
    para(doc, " Please refer to the Nondiscrimination Statement on Brightspace.",
         style="Default", bold_prefix="Non-Discrimination Statement:")
    para(doc, " Contact ODOS if you're facing food or housing insecurity; no "
              "appointment needed.", style="Default",
         bold_prefix="Basic Needs Security:")
    para(doc, " Please refer to the Emergency Preparedness on Brightspace.",
         style="Default", bold_prefix="Emergency Situation(s):")
    para(doc, " While I will try to adhere to the course schedule as much as "
              "possible, I also want to adapt to your learning pace and style. "
              "The syllabus and course plan may change during the term.",
         style="Default", bold_prefix="Subject to Change Policy:")

    # ── landscape schedule section ──────────────────────────────────────────
    end_portrait_section(doc, portrait_sectpr)
    h = para(doc, "Tentative Course Schedule · Fall 2026", space_after=0)
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)
    # D74: Davi's Word pass put a blank paragraph either side of the details
    # line and abbreviated the meeting pattern to "MWF".
    doc.add_paragraph()
    para(doc, "John Martinson Honors College, Purdue University   ·   "
              "Professor Davi Moreira   ·   Section 002, MWF, "
              "1:30–2:20 p.m., HCRS-1054   ·   "
              "August 24 – December 11, 2026", space_after=8)
    doc.add_paragraph()

    rows = build_schedule_rows()
    tbl = doc.add_table(rows=len(rows) + 1, cols=5)
    tbl.style = doc.styles["Table Grid"]
    clone_table_format(fmt.tables[-1], tbl)
    for i, head in enumerate(["Wk", "Date", "Topic", "Notebook",
                              "Assessment / Notes"]):
        cell = tbl.rows[0].cells[i]
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True
    for i, r in enumerate(rows, start=1):
        c = tbl.rows[i].cells
        c[0].text = r["wk"]
        c[1].text = r["date"]
        c[2].text = r["topic"]
        if r["nb_url"]:
            p = c[3].paragraphs[0]
            hyperlink(p, r["nb_url"], r["nb"])
        else:
            c[3].text = r["nb"]
        c[4].text = r["note"]
        if r["topic"].startswith("PURDUE") or r["note"]:
            for run in c[4].paragraphs[0].runs:
                run.bold = True
    para(doc, "Subject to change. While I will try to adhere to the course "
              "schedule as much as possible, I also want to adapt to your "
              "learning pace and style; the syllabus and course plan may change "
              "during the term. The official source of record is the course "
              "Brightspace page.")

    apply_blank_spacing(doc)
    dropped = prune_orphan_hyperlinks(doc)
    doc.save(OUT)
    patch_header_banner(OUT, BANNER, BANNER_CROP_W / 1423)
    n_nb = sum(1 for r in rows if r["nb_url"])
    print(f"✓ {OUT.relative_to(ROOT)}")
    print(f"  {len(rows)} schedule rows ({n_nb} Colab links), "
          f"{len(doc.tables)} tables, {len(doc.sections)} sections "
          f"(portrait + landscape)")
    print(f"  pruned {dropped} hyperlink relationships inherited from the "
          f"format base")


if __name__ == "__main__":
    main()
